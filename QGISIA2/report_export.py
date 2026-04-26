# -*- coding: utf-8 -*-
"""
Export de rapports SIG en PDF et DOCX.

Architecture :
    - PDF : reportlab (autonome, pure Python) ou QgsLayoutExporter si QGIS
    - DOCX : python-docx
    - Snapshot carte : QgsMapRendererCustomPainterJob (si QGIS dispo)

Le module fonctionne en 2 modes :
    1. Mode QGIS (dans le plugin) : capture la carte active + métadonnées projet
    2. Mode autonome (tests/CI) : accepte uniquement données structurées + images

Sections d'un rapport standard :
    - Page de garde (titre, date, auteur)
    - Carte principale (PNG de l'étendue projet)
    - Couches du projet (tableau : nom, type, CRS, n features)
    - Statistiques attributaires (par couche, optionnel)
    - Légende et symbologie
    - Annexes (texte libre)

Dépendances :
    pip install reportlab python-docx Pillow
"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


# ─── Détection environnement ──────────────────────────────────────────────────


def is_pdf_available() -> tuple[bool, str]:
    try:
        import reportlab  # noqa: F401
    except ImportError as e:
        return False, f"reportlab non installe : {e}. Installe via 'pip install reportlab'."
    return True, ""


def is_docx_available() -> tuple[bool, str]:
    try:
        import docx  # noqa: F401
    except ImportError as e:
        return False, f"python-docx non installe : {e}. Installe via 'pip install python-docx'."
    return True, ""


class ReportExportError(RuntimeError):
    """Levée pour toute erreur d'export rapport."""


# ─── Modèles de données ───────────────────────────────────────────────────────


@dataclass
class LayerInfo:
    name: str
    type: str  # "vector" | "raster"
    crs: str = ""
    feature_count: Optional[int] = None
    geometry_type: str = ""


@dataclass
class ReportSection:
    title: str
    body: str = ""
    bullets: list[str] = field(default_factory=list)
    table_headers: list[str] = field(default_factory=list)
    table_rows: list[list[str]] = field(default_factory=list)
    image_path: Optional[str] = None


@dataclass
class ReportConfig:
    title: str
    author: str = ""
    subtitle: str = ""
    map_image: Optional[str] = None  # PNG/JPG de la carte
    layers: list[LayerInfo] = field(default_factory=list)
    sections: list[ReportSection] = field(default_factory=list)
    footer: str = "QGISIA+ — Plugin QGIS"


@dataclass
class ExportResult:
    ok: bool
    output_path: str
    format: str  # "pdf" | "docx"
    pages_or_sections: int
    duration_s: float
    message: str


# ─── Validation ───────────────────────────────────────────────────────────────


def _validate_config(config: ReportConfig) -> None:
    if not config.title or not config.title.strip():
        raise ValueError("Le titre du rapport est obligatoire.")
    if config.map_image and not Path(config.map_image).exists():
        raise ValueError(f"Image carte introuvable : {config.map_image}")


def _validate_output_path(output_path: str, expected_ext: str) -> Path:
    if not output_path or not output_path.strip():
        raise ValueError("output_path est obligatoire.")
    p = Path(output_path)
    if p.suffix.lower() != expected_ext:
        raise ValueError(
            f"Extension attendue : {expected_ext} (recue : {p.suffix or 'aucune'})",
        )
    p.parent.mkdir(parents=True, exist_ok=True)
    return p


# ─── Export PDF (reportlab) ───────────────────────────────────────────────────


def export_pdf(config: ReportConfig, output_path: str) -> ExportResult:
    """Génère un PDF structuré via reportlab."""
    import time

    ok, reason = is_pdf_available()
    if not ok:
        raise ReportExportError(reason)

    _validate_config(config)
    out_path = _validate_output_path(output_path, ".pdf")
    start = time.time()

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=config.title,
        author=config.author or "QGISIA+",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleQGISIA", parent=styles["Title"], fontSize=22, spaceAfter=12,
    )
    h2 = styles["Heading2"]
    body = styles["BodyText"]

    story: list[Any] = []

    # Page de garde
    story.append(Paragraph(config.title, title_style))
    if config.subtitle:
        story.append(Paragraph(config.subtitle, styles["Heading3"]))
    story.append(Spacer(1, 0.5 * cm))
    meta = []
    if config.author:
        meta.append(f"<b>Auteur :</b> {config.author}")
    meta.append(f"<b>Date :</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    story.append(Paragraph("<br/>".join(meta), body))
    story.append(Spacer(1, 1 * cm))

    if config.map_image:
        try:
            story.append(Image(config.map_image, width=16 * cm, height=10 * cm))
            story.append(Spacer(1, 0.5 * cm))
        except Exception as e:  # noqa: BLE001
            logger.warning("Impossible d'inserer l'image carte : %s", e)

    # Tableau des couches
    if config.layers:
        story.append(PageBreak())
        story.append(Paragraph("Couches du projet", h2))
        rows = [["Nom", "Type", "CRS", "Geometrie", "N features"]]
        for L in config.layers:
            rows.append([
                L.name,
                L.type,
                L.crs or "—",
                L.geometry_type or "—",
                str(L.feature_count) if L.feature_count is not None else "—",
            ])
        tbl = Table(rows, repeatRows=1, colWidths=[5 * cm, 2.5 * cm, 3 * cm, 3 * cm, 2.5 * cm])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 0.5 * cm))

    # Sections
    for sec in config.sections:
        story.append(PageBreak())
        story.append(Paragraph(sec.title, h2))
        if sec.body:
            story.append(Paragraph(sec.body, body))
            story.append(Spacer(1, 0.3 * cm))
        if sec.bullets:
            for b in sec.bullets:
                story.append(Paragraph(f"• {b}", body))
            story.append(Spacer(1, 0.3 * cm))
        if sec.table_headers and sec.table_rows:
            data = [sec.table_headers] + [[str(c) for c in r] for r in sec.table_rows]
            tbl = Table(data, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 0.3 * cm))
        if sec.image_path and Path(sec.image_path).exists():
            try:
                story.append(Image(sec.image_path, width=14 * cm, height=9 * cm))
            except Exception as e:  # noqa: BLE001
                logger.warning("Image section non integree : %s", e)

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.grey)
        canvas.drawString(2 * cm, 1 * cm, config.footer)
        canvas.drawRightString(
            doc_.pagesize[0] - 2 * cm, 1 * cm, f"Page {doc_.page}",
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)

    duration = time.time() - start
    return ExportResult(
        ok=True,
        output_path=str(out_path),
        format="pdf",
        pages_or_sections=len(config.sections) + 1,
        duration_s=duration,
        message=f"PDF genere : {out_path.name} ({duration:.1f}s)",
    )


# ─── Export DOCX (python-docx) ────────────────────────────────────────────────


def export_docx(config: ReportConfig, output_path: str) -> ExportResult:
    """Génère un DOCX structuré via python-docx."""
    import time

    ok, reason = is_docx_available()
    if not ok:
        raise ReportExportError(reason)

    _validate_config(config)
    out_path = _validate_output_path(output_path, ".docx")
    start = time.time()

    from docx import Document
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    doc.core_properties.title = config.title
    if config.author:
        doc.core_properties.author = config.author

    # Page de garde
    h = doc.add_heading(config.title, level=0)
    if config.subtitle:
        doc.add_heading(config.subtitle, level=2)
    p = doc.add_paragraph()
    if config.author:
        p.add_run(f"Auteur : {config.author}\n").bold = True
    p.add_run(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}").bold = True

    if config.map_image and Path(config.map_image).exists():
        try:
            doc.add_picture(config.map_image, width=Cm(16))
        except Exception as e:  # noqa: BLE001
            logger.warning("Image carte non integree : %s", e)

    # Tableau couches
    if config.layers:
        doc.add_page_break()
        doc.add_heading("Couches du projet", level=1)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, label in enumerate(["Nom", "Type", "CRS", "Geometrie", "N features"]):
            hdr[i].text = label
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
        for L in config.layers:
            row = tbl.add_row().cells
            row[0].text = L.name
            row[1].text = L.type
            row[2].text = L.crs or "—"
            row[3].text = L.geometry_type or "—"
            row[4].text = str(L.feature_count) if L.feature_count is not None else "—"

    # Sections
    for sec in config.sections:
        doc.add_page_break()
        doc.add_heading(sec.title, level=1)
        if sec.body:
            doc.add_paragraph(sec.body)
        for b in sec.bullets:
            doc.add_paragraph(b, style="List Bullet")
        if sec.table_headers and sec.table_rows:
            tbl = doc.add_table(rows=1, cols=len(sec.table_headers))
            tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(sec.table_headers):
                tbl.rows[0].cells[i].text = str(h)
            for r in sec.table_rows:
                cells = tbl.add_row().cells
                for i, val in enumerate(r):
                    if i < len(cells):
                        cells[i].text = str(val)
        if sec.image_path and Path(sec.image_path).exists():
            try:
                doc.add_picture(sec.image_path, width=Cm(14))
            except Exception as e:  # noqa: BLE001
                logger.warning("Image section non integree : %s", e)

    # Footer
    section = doc.sections[0]
    footer_p = section.footer.paragraphs[0]
    footer_run = footer_p.add_run(config.footer)
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(out_path))
    duration = time.time() - start
    return ExportResult(
        ok=True,
        output_path=str(out_path),
        format="docx",
        pages_or_sections=len(config.sections) + 1,
        duration_s=duration,
        message=f"DOCX genere : {out_path.name} ({duration:.1f}s)",
    )


# ─── API publique ─────────────────────────────────────────────────────────────


def export_report(
    config: ReportConfig,
    output_path: str,
    format: str = "pdf",
) -> ExportResult:
    """
    Export d'un rapport SIG en PDF ou DOCX.

    Args:
        config: ReportConfig avec titre, couches, sections.
        output_path: chemin de sortie (.pdf ou .docx).
        format: "pdf" (defaut) ou "docx".

    Raises:
        ValueError pour params invalides.
        ReportExportError si dependances manquantes.
    """
    fmt = format.lower().strip()
    if fmt == "pdf":
        return export_pdf(config, output_path)
    if fmt == "docx":
        return export_docx(config, output_path)
    raise ValueError(f"Format non supporte : {format}. Choisis 'pdf' ou 'docx'.")
